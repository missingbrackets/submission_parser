# ============================================================
#  FILE PARSER
#  Extracts raw text from PDF, Excel, Word, and .msg files
#  in a submission folder ready to send to Claude.
# ============================================================

import os
import io
import traceback
from pathlib import Path
from typing import Tuple

# ── PDF ──────────────────────────────────────────────────────
def extract_pdf(filepath: str) -> Tuple[str, str]:
    """Returns (text, status_message)"""
    try:
        import pdfplumber
        text_parts = []
        with pdfplumber.open(filepath) as pdf:
            for i, page in enumerate(pdf.pages):
                t = page.extract_text()
                if t:
                    text_parts.append(f"[Page {i+1}]\n{t}")
                # Also extract any tables
                tables = page.extract_tables()
                for tbl in tables:
                    rows = []
                    for row in tbl:
                        rows.append(" | ".join([str(c) if c else "" for c in row]))
                    if rows:
                        text_parts.append("[TABLE]\n" + "\n".join(rows))
        full_text = "\n\n".join(text_parts)
        if not full_text.strip():
            return "", "PDF parsed but no text found (may be scanned/image-based)"
        return full_text, f"PDF extracted: {len(full_text)} chars, {len(pdf.pages)} pages"
    except ImportError:
        return "", "pdfplumber not installed — run: pip install pdfplumber"
    except Exception as e:
        return "", f"PDF extraction failed: {str(e)}"


# ── EXCEL ────────────────────────────────────────────────────
def extract_excel(filepath: str) -> Tuple[str, str]:
    """Returns (text, status_message)"""
    try:
        import openpyxl
        wb = openpyxl.load_workbook(filepath, data_only=True)
        text_parts = []
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            rows = []
            for row in ws.iter_rows(values_only=True):
                # Skip completely empty rows
                if any(cell is not None for cell in row):
                    rows.append(" | ".join([str(c) if c is not None else "" for c in row]))
            if rows:
                text_parts.append(f"[Sheet: {sheet_name}]\n" + "\n".join(rows))
        full_text = "\n\n".join(text_parts)
        if not full_text.strip():
            return "", "Excel parsed but all sheets appear empty"
        return full_text, f"Excel extracted: {len(wb.sheetnames)} sheet(s), {len(full_text)} chars"
    except ImportError:
        return "", "openpyxl not installed — run: pip install openpyxl"
    except Exception as e:
        # Try xlrd as fallback for .xls
        try:
            import xlrd
            wb = xlrd.open_workbook(filepath)
            text_parts = []
            for sheet in wb.sheets():
                rows = []
                for r in range(sheet.nrows):
                    row = [str(sheet.cell_value(r, c)) for c in range(sheet.ncols)]
                    if any(v.strip() for v in row):
                        rows.append(" | ".join(row))
                if rows:
                    text_parts.append(f"[Sheet: {sheet.name}]\n" + "\n".join(rows))
            full_text = "\n\n".join(text_parts)
            return full_text, f"Excel (.xls) extracted via xlrd: {len(full_text)} chars"
        except Exception as e2:
            return "", f"Excel extraction failed: {str(e)} / {str(e2)}"


# ── WORD ─────────────────────────────────────────────────────
def extract_word(filepath: str) -> Tuple[str, str]:
    """Returns (text, status_message)"""
    try:
        from docx import Document
        doc = Document(filepath)
        parts = []

        # Body paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text)

        # Tables
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                if any(cells):
                    parts.append(" | ".join(cells))

        full_text = "\n".join(parts)
        if not full_text.strip():
            return "", "Word doc parsed but no text found"
        return full_text, f"Word doc extracted: {len(full_text)} chars"
    except ImportError:
        return "", "python-docx not installed — run: pip install python-docx"
    except Exception as e:
        return "", f"Word extraction failed: {str(e)}"


# ── MSG (Outlook email) ───────────────────────────────────────
def extract_msg(filepath: str) -> Tuple[str, str]:
    """Returns (text, status_message)"""
    try:
        import extract_msg as em
        msg = em.Message(filepath)
        parts = []
        if msg.subject:
            parts.append(f"Subject: {msg.subject}")
        if msg.sender:
            parts.append(f"From: {msg.sender}")
        if msg.date:
            parts.append(f"Date: {msg.date}")
        if msg.body:
            parts.append(f"\n[Email Body]\n{msg.body}")
        full_text = "\n".join(parts)
        return full_text, f".msg extracted: {len(full_text)} chars"
    except ImportError:
        return "", "extract-msg not installed — run: pip install extract-msg"
    except Exception as e:
        return "", f".msg extraction failed: {str(e)}"


# ── PLAIN TEXT ────────────────────────────────────────────────
def extract_text(filepath: str) -> Tuple[str, str]:
    try:
        with open(filepath, "r", encoding="utf-8", errors="replace") as f:
            text = f.read()
        return text, f"Text file read: {len(text)} chars"
    except Exception as e:
        return "", f"Text read failed: {str(e)}"


# ── DISPATCHER ────────────────────────────────────────────────
EXTENSION_MAP = {
    ".pdf":  extract_pdf,
    ".xlsx": extract_excel,
    ".xls":  extract_excel,
    ".xlsm": extract_excel,
    ".docx": extract_word,
    ".doc":  extract_word,
    ".msg":  extract_msg,
    ".txt":  extract_text,
    ".csv":  extract_text,
}

def extract_file(filepath: str) -> Tuple[str, str]:
    """
    Dispatch to correct extractor based on file extension.
    Returns (extracted_text, status_message)
    """
    ext = Path(filepath).suffix.lower()
    fn = EXTENSION_MAP.get(ext)
    if fn is None:
        return "", f"Unsupported file type: {ext}"
    return fn(filepath)


def extract_folder(folder_path: str, subfolder: str = None) -> dict:
    """
    Extract all supported files from a folder (or subfolder).
    Returns dict: { filename: { "text": str, "status": str, "path": str } }
    """
    results = {}
    target = os.path.join(folder_path, subfolder) if subfolder else folder_path

    if not os.path.exists(target):
        return results

    for fname in sorted(os.listdir(target)):
        fpath = os.path.join(target, fname)
        if not os.path.isfile(fpath):
            continue
        ext = Path(fname).suffix.lower()
        if ext not in EXTENSION_MAP:
            continue
        text, status = extract_file(fpath)
        results[fname] = {
            "text": text,
            "status": status,
            "path": fpath,
            "size_kb": round(os.path.getsize(fpath) / 1024, 1),
        }

    return results
